"""Репозиторий для сохранения вопросов/ответов в docx-файл интервью."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.table import Table

from app.features.chat.domain.interview_docx import load_interview_qa

logger = logging.getLogger(__name__)

# Ширина колонок таблицы из оригинального файла (в EMU).
# Используются при создании нового файла; при дописывании в существующий
# таблица уже имеет правильные размеры.
_COLUMN_WIDTHS_EMU = (535305, 1947545, 6991350)
_FONT_NAME = "Times New Roman"
_FONT_SIZE_PT = 14.0  # 177800 EMU == 14pt


def _find_table(doc: Document) -> Table:
    """
    Находит единственную таблицу с вопросами в документе.
    :param doc: открытый Document
    :return: объект Table
    :raises ValueError: если таблица не найдена
    """
    if not doc.tables:
        raise ValueError("В документе отсутствует таблица с вопросами")
    return doc.tables[0]


def _matches_question(existing_question: str, new_question: str) -> bool:
    """
    Сравнивает два вопроса на совпадение (нормализованное сравнение).
    Убирает лишние пробелы, игнорирует регистр и конечные знаки препинания.
    :param existing_question: вопрос из файла
    :param new_question: новый вопрос для сохранения
    :return: True, если вопросы совпадают
    """
    import re

    def normalize(text: str) -> str:
        text = re.sub(r"\s+", " ", text or "").strip()
        text = text.rstrip("?.!")
        return text.lower()

    return normalize(existing_question) == normalize(new_question)


def question_exists(path: Path, question: str) -> bool:
    """
    Проверяет, есть ли уже такой вопрос в docx-файле.
    :param path: путь к docx
    :param question: текст вопроса для поиска
    :return: True, если вопрос уже присутствует
    """
    if not path.exists():
        return False
    try:
        existing = load_interview_qa(path)
    except Exception:
        return False
    return any(_matches_question(item.question, question) for item in existing)


def _max_question_number(path: Path) -> int:
    """
    Возвращает максимальный номер вопроса в файле.
    Если файл не существует или пуст — 0.
    """
    if not path.exists():
        return 0
    try:
        items = load_interview_qa(path)
    except Exception:
        return 0
    return max((item.number for item in items), default=0)


def _set_cell_text(cell: object, text: str, *, bold: bool = False) -> None:
    """
        Записывает текст в первый параграф ячейки, устанавливая шрифт.
        Удаляет существующие раны
    , чтобы не было пустых runs,
        которые ломают чтение python-docx после сохранения.
        :param cell: ячейка таблицы (docx _Cell)
        :param text: текст для записи
        :param bold: сделать текст жирным
    """
    from docx.text.paragraph import Paragraph as _Para

    para: _Para = cell.paragraphs[0]  # type: ignore[attr-defined]
    # Удаляем все существующие раны, чтобы не было пустых
    for run in para.runs:
        run._element.getparent().remove(run._element)
    run = para.add_run(text)
    run.font.name = _FONT_NAME
    run.font.size = Pt(_FONT_SIZE_PT)
    run.bold = bold


def _clean_text(text: str) -> str:
    """
    Очищает текст от лишних символов форматирования.
    Убирает markdown-разметку, эмодзи, лишние пробелы и пустые строки.
    :param text: исходный текст
    :return: очищенный текст
    """
    import re

    if not text:
        return ""

    # Убираем markdown-заголовки (#, ##, ###)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)

    # Убираем markdown-жирный/курсив (**text**, *text*, __text__, _text_)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"(?<![\w])_(.+?)_(?!\w)", r"\1", text)

    # Убираем markdown-кодовые блоки (```)
    text = re.sub(r"```[a-zA-Z]*\n?", "", text)

    # Убираем инлайн-код (`code`)
    text = re.sub(r"`(.+?)`", r"\1", text)

    # Убираем эмодзи
    emoji_pattern = re.compile(
        "["
        "\U0001f600-\U0001f64f"  # emoticons
        "\U0001f300-\U0001f5ff"  # symbols & pictographs
        "\U0001f680-\U0001f6ff"  # transport & map
        "\U0001f1e0-\U0001f1ff"  # flags
        "\U00002702-\U000027b0"
        "\U000024c2-\U0001f251"
        "\U0001f926-\U0001f937"
        "\U00010000-\U0010ffff"
        "\u2640-\u2642"
        "\u2600-\u2b55"
        "\u200d"
        "\u23cf"
        "\u23e9"
        "\u231a"
        "\ufe0f"
        "\u3030"
        "]+",
        flags=re.UNICODE,
    )
    text = emoji_pattern.sub("", text)

    # Убираем спецсимволы markdown (тире, звёздки списков)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)

    # Убираем строки с ключевыми словами (Оценка, Источники, Дополнение)
    # Также убираем строки вида "Источники: ответы №..." и "Источник: ответ №..."
    lines = text.split("\n")
    cleaned_lines = []
    skip_mode = False
    for line in lines:
        stripped = line.strip()
        # Если строка начинается с ключевого слова — пропускаем её и всё после до пустой строки
        if re.match(r"^(Оценка|Источники?|Дополнение|Дополнения)\b", stripped, re.IGNORECASE):
            skip_mode = True
            continue
        # Убираем строки "Источник: ответ №N" и "Источники: ответы №N" (в любом месте)
        if re.match(r"^Источники?\s*:\s*ответ", stripped, re.IGNORECASE):
            continue
        if skip_mode:
            if stripped == "":
                skip_mode = False
            continue
        cleaned_lines.append(line)
    text = "\n".join(cleaned_lines)

    # Убираем оставшиеся пустые строки
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Убираем пустые строки (оставляем одну)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Убираем лишние пробелы в начале/конце строк
    lines = [line.strip() for line in text.split("\n")]
    text = "\n".join(lines)

    return text.strip()


def _add_row(table: Table, number: int, question: str, answer: str) -> None:
    """
    Добавляет строку с вопросом/ответом в конец таблицы.
    Форматирование повторяет стиль оригинального файла:
    - Times New Roman, 14pt
    - номер жирный
    - вопрос жирный
    - первая строка ответа жирная (заголовок), остальное — обычным текстом
    :param table: таблица docx
    :param number: номер вопроса
    :param question: текст вопроса
    :param answer: текст ответа
    """
    row = table.add_row()

    # Устанавливаем ширину колонок для новой строки
    for i, cell in enumerate(row.cells):
        cell.width = _COLUMN_WIDTHS_EMU[i]  # type: ignore[attr-defined]

    # Колонка 0: номер (жирный)
    _set_cell_text(row.cells[0], str(number), bold=True)

    # Колонка 1: вопрос (жирный)
    _set_cell_text(row.cells[1], question, bold=True)

    # Колонка 2: ответ
    # Первая строка ответа — жирная (краткий заголовок/тезис),
    # последующий текст — обычный.
    cell_a = row.cells[2]
    answer_lines = answer.split("\n") if answer else [""]
    for line_index, line in enumerate(answer_lines):
        if line_index == 0:
            p_line = cell_a.paragraphs[0]
        else:
            p_line = cell_a.add_paragraph()
        # Удаляем пустые раны из параграфа
        for run in p_line.runs:
            run._element.getparent().remove(run._element)
        run_line = p_line.add_run(line)
        run_line.font.name = _FONT_NAME
        run_line.font.size = Pt(_FONT_SIZE_PT)
        # Первая строка жирная, если она выглядит как заголовок
        # (заканчивается двоеточием или это единственная строка в списке)
        run_line.bold = line_index == 0 and (line.rstrip().endswith(":") or len(answer_lines) == 1)


def save_question_answer(
    path: Path,
    question: str,
    answer: str,
    *,
    force: bool = False,
) -> dict[str, object]:
    """
    Сохраняет вопрос/ответ в конец docx-файла, если такого вопроса ещё нет.
    Файл должен существовать и содержать таблицу с 3 колонками.
    :param path: путь к docx-файлу
    :param question: текст вопроса
    :param answer: текст ответа
    :param force: если True, сохраняет даже при дубликате (для тестов)
    :return: словарь с результатом операции
    :raises FileNotFoundError: если файл не найден
    """
    if not path.exists():
        raise FileNotFoundError(f"Файл не найден: {path}")

    if not force and question_exists(path, question):
        logger.info("Вопрос уже существует в %s, пропускаем сохранение", path)
        return {"status": "skipped", "reason": "already_exists"}

    doc = Document(str(path))
    table = _find_table(doc)

    next_number = _max_question_number(path) + 1
    clean_question = _clean_text(question)
    clean_answer = _clean_text(answer)
    _add_row(table, next_number, clean_question, clean_answer)
    doc.save(str(path))

    logger.info(
        "Сохранён вопрос №%d в %s: %s",
        next_number,
        path.name,
        question[:80],
    )
    return {"status": "saved", "number": next_number}
