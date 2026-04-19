"""
Парсинг ``.docx`` файла с вопросами/ответами для подготовки к собеседованиям.

Модуль старается аккуратно извлечь пары "вопрос/ответ" даже если документ размечен
нестрого. Номер вопроса извлекается из заголовка (например ``12. ...``).
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document

QUESTION_PATTERN = re.compile(
    r"^\s*(?:вопрос\s*)?(\d{1,4})\s*[.)\-:]\s*(.+)?$", re.IGNORECASE
)
ANSWER_MARKER_PATTERN = re.compile(r"^\s*ответ\s*[:\-]?\s*(.*)$", re.IGNORECASE)


@dataclass(slots=True)
class InterviewQA:
    """Одна запись "вопрос/ответ" из docx."""

    number: int
    question: str
    answer: str

    @property
    def as_document(self) -> str:
        """Текстовый документ для RAG-индексации."""
        return f"Вопрос №{self.number}\n{self.question}\n\nОтвет №{self.number}\n{self.answer}"


def _normalize_paragraphs(path: Path) -> list[str]:
    doc = Document(str(path))
    out: list[str] = []
    for p in doc.paragraphs:
        text = re.sub(r"\s+", " ", p.text or "").strip()
        if text:
            out.append(text)
    return out


def _load_from_tables(path: Path) -> list[InterviewQA]:
    """
    Извлекает пары вопрос/ответ из табличной структуры.

    Ожидается формат строк: [номер, вопрос, ответ].
    """
    doc = Document(str(path))
    out: list[InterviewQA] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [re.sub(r"\s+", " ", c.text or "").strip() for c in row.cells]
            if len(cells) < 3:
                continue
            number_str = cells[0]
            if not number_str.isdigit():
                # Обычно первая строка — заголовок колонок.
                continue
            number = int(number_str)
            question = cells[1].strip()
            answer = cells[2].strip()
            if question and answer:
                out.append(InterviewQA(number=number, question=question, answer=answer))
    return out


def load_interview_qa(path: Path) -> list[InterviewQA]:
    """
    Читает ``.docx`` и извлекает список вопросов и ответов.

    Поддерживаются маркеры вида:
    - ``12. Что такое GIL``
    - ``Вопрос 12: ...``
    - ``Ответ: ...`` (опционально)
    """
    # 1) Попробуем явную табличную структуру (наиболее частый формат рабочих шпаргалок).
    table_items = _load_from_tables(path)
    if table_items:
        return table_items

    # 2) Fallback: парсинг последовательности абзацев.
    paragraphs = _normalize_paragraphs(path)
    if not paragraphs:
        return []

    items: list[InterviewQA] = []
    cur_number: int | None = None
    cur_question = ""
    cur_answer_parts: list[str] = []
    collecting_answer = False

    def flush() -> None:
        nonlocal cur_number, cur_question, cur_answer_parts, collecting_answer
        if cur_number is None:
            return
        answer = " ".join(cur_answer_parts).strip()
        items.append(
            InterviewQA(
                number=cur_number,
                question=cur_question.strip() or f"Вопрос №{cur_number}",
                answer=answer,
            )
        )
        cur_number = None
        cur_question = ""
        cur_answer_parts = []
        collecting_answer = False

    for line in paragraphs:
        q_match = QUESTION_PATTERN.match(line)
        if q_match:
            flush()
            cur_number = int(q_match.group(1))
            cur_question = (q_match.group(2) or "").strip()
            continue

        if cur_number is None:
            continue

        a_match = ANSWER_MARKER_PATTERN.match(line)
        if a_match:
            collecting_answer = True
            marker_tail = a_match.group(1).strip()
            if marker_tail:
                cur_answer_parts.append(marker_tail)
            continue

        if collecting_answer:
            cur_answer_parts.append(line)
        elif not cur_question:
            cur_question = line
        else:
            # Иногда ответы идут сразу после вопроса без "Ответ:"
            cur_answer_parts.append(line)

    flush()
    return items
