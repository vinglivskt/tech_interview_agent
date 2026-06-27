"""Тесты для репозитория сохранения вопросов/ответов в docx."""

from __future__ import annotations

import shutil
from pathlib import Path

import pytest

from app.features.chat.domain.docx_repository import (
    question_exists,
    save_question_answer,
)

FIXTURES_DIR = Path(__file__).resolve().parent.parent / "fixtures"
SRC_FIXTURE = FIXTURES_DIR / "test_interview_questions.docx"


@pytest.fixture
def tmp_docx(tmp_path: Path) -> Path:
    """
    Создаёт свежую копию тестового docx-файла для каждого теста.
    Работаем с копией, чтобы не мутировать эталон.
    """
    dst = tmp_path / "interview_test.docx"
    shutil.copy2(SRC_FIXTURE, dst)
    return dst


def _row_count(path: Path) -> int:
    from docx import Document

    doc = Document(str(path))
    return len(doc.tables[0].rows)


def _last_row(path: Path) -> list[str]:
    from docx import Document

    doc = Document(str(path))
    row = doc.tables[0].rows[-1]
    return [c.text.strip() for c in row.cells]


def test_save_new_question_appends_row(tmp_docx: Path) -> None:
    """Новый вопрос добавляется в конец таблицы с правильным номером."""
    initial_rows = _row_count(tmp_docx)
    assert initial_rows == 4  # header + 3 sample rows

    result = save_question_answer(
        tmp_docx,
        "Что такое декоратор в Python?",
        "Декоратор — это функция, которая принимает другую функцию и расширяет её поведение без изменения исходного кода.\n\nСинтаксис:\n@decorator\ndef func(): ...",
    )

    assert result["status"] == "saved"
    assert result["number"] == 4  # next after 1, 2, 3

    # Row was appended
    assert _row_count(tmp_docx) == initial_rows + 1

    # Verify content of the last row
    last = _last_row(tmp_docx)
    assert last[0] == "4"
    assert "декоратор" in last[1].lower()
    assert "Декоратор" in last[2]


def test_save_duplicate_question_is_skipped(tmp_docx: Path) -> None:
    """Если вопрос уже есть в файле — дубликат не сохраняется."""
    initial_rows = _row_count(tmp_docx)

    # Try to save a question that already exists (row 2 in fixture)
    result = save_question_answer(
        tmp_docx,
        "Механизм очистки памяти в Python",
        "Какой-то ответ",
    )

    assert result["status"] == "skipped"
    assert result["reason"] == "already_exists"
    assert _row_count(tmp_docx) == initial_rows  # no new row


def test_save_duplicate_with_different_punctuation_is_skipped(tmp_docx: Path) -> None:
    """Дубликат определяется независимо от регистра и конечных знаков препинания."""
    initial_rows = _row_count(tmp_docx)

    result = save_question_answer(
        tmp_docx,
        "механизм очистки памяти в python???",
        "Ответ",
    )

    assert result["status"] == "skipped"
    assert _row_count(tmp_docx) == initial_rows


def test_save_multiple_new_questions_increment_numbers(tmp_docx: Path) -> None:
    """Несколько новых вопросов получают последовательные номера."""
    r1 = save_question_answer(tmp_docx, "Вопрос А?", "Ответ А")
    r2 = save_question_answer(tmp_docx, "Вопрос Б?", "Ответ Б")
    r3 = save_question_answer(tmp_docx, "Вопрос В?", "Ответ В")

    assert r1["number"] == 4
    assert r2["number"] == 5
    assert r3["number"] == 6

    assert _row_count(tmp_docx) == 7  # header + 3 sample + 3 new


def test_question_exists_returns_true_for_existing(tmp_docx: Path) -> None:
    """question_exists находит существующий вопрос."""
    assert question_exists(tmp_docx, "Mock объект тестирования. Что это?") is True


def test_question_exists_returns_false_for_new(tmp_docx: Path) -> None:
    """question_exists возвращает False для нового вопроса."""
    assert question_exists(tmp_docx, "Что такое FastAPI?") is False


def test_save_to_nonexistent_file_raises(tmp_path: Path) -> None:
    """Сохранение в несуществующий файл выбрасывает FileNotFoundError."""
    with pytest.raises(FileNotFoundError):
        save_question_answer(tmp_path / "missing.docx", "Вопрос?", "Ответ")


def test_save_preserves_existing_content(tmp_docx: Path) -> None:
    """После сохранения нового вопроса все предыдущие строки остаются нетронутыми."""
    from docx import Document

    # Read original content
    doc_before = Document(str(tmp_docx))
    rows_before = [[c.text for c in row.cells] for row in doc_before.tables[0].rows]

    save_question_answer(tmp_docx, "Совершенно новый вопрос", "Совершенно новый ответ")

    # Verify all previous rows are intact
    doc_after = Document(str(tmp_docx))
    rows_after = [[c.text for c in row.cells] for row in doc_after.tables[0].rows[: len(rows_before)]]

    assert rows_before == rows_after


def test_save_preserves_formatting(tmp_docx: Path) -> None:
    """Новая строка структурно совпадает с существующими (жирность номера/вопроса)."""
    from docx import Document

    save_question_answer(
        tmp_docx,
        "Тестовый вопрос форматирования",
        "Заголовок ответа:\nПодробное объяснение второго параграфа.",
    )

    doc = Document(str(tmp_docx))
    new_row = doc.tables[0].rows[-1]

    # Number cell should be bold (like existing rows)
    num_run = new_row.cells[0].paragraphs[0].runs[0]
    assert num_run.bold is True

    # Question cell should be bold (like existing rows)
    q_run = new_row.cells[1].paragraphs[0].runs[0]
    assert q_run.bold is True

    # Answer first line should be bold (ends with colon)
    a_first_run = new_row.cells[2].paragraphs[0].runs[0]
    assert a_first_run.bold is True

    # Second line of answer should NOT be bold
    a_second_run = new_row.cells[2].paragraphs[1].runs[0]
    assert a_second_run.bold is not True
